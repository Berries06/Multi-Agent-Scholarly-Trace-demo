from __future__ import annotations

import argparse
import math
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "submission_materials"
ASSET_DIR = OUTPUT_DIR / "assets"
DOCX_NAME = "研海寻踪：基于多智能体博弈推理的科研知识图谱发现系统+作品书（预审核初稿）.docx"

TEAL = "087F78"
TEAL_DARK = "075E59"
NAVY = "17324D"
GOLD = "D79A2B"
INK = "243238"
MUTED = "66777C"
LIGHT_TEAL = "E7F3F1"
LIGHT_BLUE = "EAF0F6"
LIGHT_GOLD = "FBF3DF"
LIGHT_GRAY = "F3F5F4"
RED = "A54737"
WHITE = "FFFFFF"
BLACK = "000000"

BODY_FONT_CN = "宋体"
BODY_FONT_EN = "Times New Roman"
HEADING_FONT_CN = "黑体"
FONT_REGULAR = Path("C:/Windows/Fonts/simsun.ttc")
FONT_BOLD = Path("C:/Windows/Fonts/simhei.ttf")


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color: str = "B7C5C2", size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_run_font(
    run,
    *,
    cn: str = BODY_FONT_CN,
    en: str = BODY_FONT_EN,
    size: float = 12,
    bold: bool | None = None,
    color: str = INK,
    italic: bool | None = None,
) -> None:
    run.font.name = en
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cn)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), en)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), en)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph(
    paragraph,
    *,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    before: float = 0,
    after: float = 6,
    line: float = 1.5,
    first_line: bool = False,
    keep_with_next: bool = False,
) -> None:
    pf = paragraph.paragraph_format
    pf.alignment = alignment
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if first_line:
        pf.first_line_indent = Pt(24)
    pf.keep_with_next = keep_with_next
    pf.widow_control = True


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, first_line=bold_lead is None)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        tail = p.add_run(text[len(bold_lead) :])
        set_run_font(tail)
    else:
        run = p.add_run(text)
        set_run_font(run)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.74 + level * 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.37)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    set_run_font(run, size=11.5)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.9)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    set_run_font(run, size=11.5)


def add_numbered_list(doc: Document, items: Iterable[str]) -> None:
    """Add a real numbered list whose numbering restarts at 1."""
    numbering = doc.part.numbering_part.element
    base_num_id = int(doc.styles["List Number"].element.pPr.numPr.numId.val)
    base_num = numbering.xpath(f'./w:num[@w:numId="{base_num_id}"]')[0]
    abstract_num_id = int(base_num.xpath("./w:abstractNumId")[0].get(qn("w:val")))
    existing_ids = [int(node.get(qn("w:numId"))) for node in numbering.xpath("./w:num")]
    num_id = max(existing_ids, default=0) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)

    for text in items:
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_element = OxmlElement("w:numId")
        num_id_element.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_element])
        p_pr.append(num_pr)
        p.paragraph_format.left_indent = Cm(0.9)
        p.paragraph_format.first_line_indent = Cm(-0.45)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.35
        run = p.add_run(text)
        set_run_font(run, size=11.5)


def add_heading(doc: Document, text: str, level: int = 1, *, page_break: bool | None = None) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    if page_break if page_break is not None else level == 1:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt({1: 16, 2: 12, 3: 8}[level])
    p.paragraph_format.space_after = Pt({1: 9, 2: 6, 3: 4}[level])
    run = p.add_run(text)
    set_run_font(
        run,
        cn=HEADING_FONT_CN,
        en="Arial",
        size={1: 18, 2: 14, 3: 12}[level],
        bold=True,
        color={1: TEAL_DARK, 2: NAVY, 3: NAVY}[level],
    )


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=8, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=9.5, color=MUTED)


def add_note(doc: Document, label: str, text: str, fill: str = LIGHT_GOLD, color: str = NAVY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color=fill, size=4)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=0, line=1.3)
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, cn=HEADING_FONT_CN, en="Arial", size=11, bold=True, color=color)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=color)


def add_table(
    doc: Document,
    headers: list[str],
    rows: Iterable[Iterable[str]],
    widths_dxa: list[int],
    *,
    header_fill: str = TEAL_DARK,
    header_color: str = WHITE,
    font_size: float = 9.5,
) -> None:
    rows = [list(row) for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.15)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_run_font(run, cn=HEADING_FONT_CN, en="Arial", size=font_size, bold=True, color=header_color)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            if row_index % 2 == 1:
                set_cell_shading(cells[idx], LIGHT_GRAY)
            p = cells[idx].paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, alignment=align, after=0, line=1.2)
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=INK)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def set_section_furniture(section, *, first_page: bool = False) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = first_page
    if not first_page:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("研海寻踪：基于多智能体博弈推理的科研知识图谱发现系统")
        set_run_font(run, size=8.5, color=MUTED)
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "B7C5C2")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
        add_page_number(section.footer.paragraphs[0])


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT_EN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT_CN)
    normal.font.size = Pt(12)
    normal.font.color.rgb = rgb(INK)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level, size, color in ((1, 18, TEAL_DARK), (2, 14, NAVY), (3, 12, NAVY)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = BODY_FONT_EN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT_CN)
        style.font.size = Pt(11.5)


def load_font(path: Path, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(path), size=size)


def draw_centered(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill: str, spacing: int = 6) -> None:
    left, top, right, bottom = box
    lines = text.split("\n")
    heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + spacing * max(0, len(lines) - 1)
    y = top + (bottom - top - total_h) / 2
    for line, width, height in zip(lines, widths, heights):
        draw.text((left + (right - left - width) / 2, y), line, font=font, fill=fill)
        y += height + spacing


def build_architecture_figure(path: Path) -> None:
    image = Image.new("RGB", (1600, 920), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(FONT_BOLD, 48)
    layer_font = load_font(FONT_BOLD, 31)
    box_font = load_font(FONT_REGULAR, 28)
    small_font = load_font(FONT_REGULAR, 23)
    draw.text((70, 38), "研海寻踪端到端系统架构", font=title_font, fill=f"#{NAVY}")
    draw.text((70, 100), "证据优先建图 · 异质三智能体裁决 · 意图驱动图检索", font=small_font, fill=f"#{MUTED}")

    layers = [
        ("输入与画像", 160, ["论文/知识卡", "领域文献元数据", "学习者画像", "研究查询"]),
        ("论文到图谱", 315, ["结构解析", "知识抽取 Agent", "实体规范化", "证据跨度图"]),
        ("可信裁决", 470, ["提出者\n候选组织", "批判者\n反证/约束", "裁判\n接收/复核/拒绝"]),
        ("检索与发现", 625, ["意图感知 Agent", "广度检索\n找论文", "深度检索\n做分析", "混合检索\n找 Idea"]),
        ("输出与工程", 780, ["技术脉络", "争议/空白", "个性化资源", "超时·幂等·熔断·审计"]),
    ]
    layer_colors = [LIGHT_BLUE, LIGHT_TEAL, LIGHT_GOLD, LIGHT_TEAL, LIGHT_GRAY]
    box_fills = ["F7FAFC", "F5FBFA", "FFF9EC", "F5FBFA", "F7F8F7"]
    for idx, (label, y, items) in enumerate(layers):
        draw.rounded_rectangle((70, y, 1530, y + 120), radius=24, fill=f"#{layer_colors[idx]}", outline=f"#{TEAL}", width=3)
        draw.rounded_rectangle((85, y + 18, 265, y + 102), radius=18, fill=f"#{TEAL_DARK}")
        draw_centered(draw, (85, y + 18, 265, y + 102), label, layer_font, "white")
        count = len(items)
        start_x = 300
        gap = 24
        width = int((1210 - gap * (count - 1)) / count)
        for item_idx, item in enumerate(items):
            x = start_x + item_idx * (width + gap)
            draw.rounded_rectangle((x, y + 18, x + width, y + 102), radius=16, fill=f"#{box_fills[idx]}", outline=f"#{NAVY}", width=2)
            draw_centered(draw, (x + 8, y + 18, x + width - 8, y + 102), item, box_font, f"#{INK}", spacing=3)
        if idx < len(layers) - 1:
            cx = 800
            draw.line((cx, y + 120, cx, y + 150), fill=f"#{TEAL_DARK}", width=7)
            draw.polygon([(cx - 13, y + 142), (cx + 13, y + 142), (cx, y + 158)], fill=f"#{TEAL_DARK}")
    image.save(path, quality=95)


def build_ablation_figure(path: Path) -> None:
    image = Image.new("RGB", (1500, 820), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(FONT_BOLD, 46)
    label_font = load_font(FONT_REGULAR, 27)
    value_font = load_font(FONT_BOLD, 27)
    small_font = load_font(FONT_REGULAR, 22)
    draw.text((60, 35), "同一候选池下的决策机制对比（24 条压力命题）", font=title_font, fill=f"#{NAVY}")
    draw.text((60, 95), "接收精确率越高越好；不支持命题接收率（UAR）越低越好", font=small_font, fill=f"#{MUTED}")
    variants = [
        ("普通规则程序", 0.500, 1.000),
        ("单次判定", 0.611, 0.636),
        ("同质三路投票", 0.611, 0.636),
        ("提出-批判-裁判", 0.846, 0.182),
    ]
    left, top, right, bottom = 130, 170, 1430, 680
    for tick in range(0, 101, 20):
        y = bottom - (bottom - top) * tick / 100
        draw.line((left, y, right, y), fill="#D9E0DE", width=2)
        draw.text((55, y - 14), f"{tick}%", font=small_font, fill=f"#{MUTED}")
    group_w = (right - left) / len(variants)
    bar_w = 95
    for idx, (label, precision, uar) in enumerate(variants):
        cx = left + group_w * (idx + 0.5)
        for offset, value, fill, metric in (
            (-bar_w - 8, precision, f"#{TEAL_DARK}", "精确率"),
            (8, uar, f"#{GOLD}", "UAR"),
        ):
            x0 = cx + offset
            x1 = x0 + bar_w
            y0 = bottom - (bottom - top) * value
            draw.rounded_rectangle((x0, y0, x1, bottom), radius=8, fill=fill)
            draw.text((x0 + 4, y0 - 37), f"{value * 100:.1f}%", font=value_font, fill=f"#{INK}")
        bbox = draw.textbbox((0, 0), label, font=label_font)
        draw.text((cx - (bbox[2] - bbox[0]) / 2, bottom + 25), label, font=label_font, fill=f"#{INK}")
    draw.rounded_rectangle((1010, 105, 1195, 145), radius=10, fill=f"#{TEAL_DARK}")
    draw.text((1210, 109), "接收精确率", font=small_font, fill=f"#{INK}")
    draw.rounded_rectangle((1300, 105, 1365, 145), radius=10, fill=f"#{GOLD}")
    draw.text((1380, 109), "UAR", font=small_font, fill=f"#{INK}")
    draw.text((110, 748), "注意：这是开发可见的小样本压力集，不代表公开基准上的端到端抽取性能。", font=small_font, fill=f"#{RED}")
    image.save(path, quality=95)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    set_section_furniture(section, first_page=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("“挑战杯”全国大学生课外学术科技作品竞赛\n“揭榜挂帅”专项赛")
    set_run_font(run, cn=HEADING_FONT_CN, en="Arial", size=17, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("成果汇报书")
    set_run_font(run, cn=HEADING_FONT_CN, en="Arial", size=28, bold=True, color=TEAL_DARK)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rule.add_run("━━━━━━━━━━━━━━━━━━━━━━━━")
    set_run_font(r, cn=HEADING_FONT_CN, en="Arial", size=13, color=GOLD)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(36)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run("研海寻踪\n")
    set_run_font(r1, cn=HEADING_FONT_CN, en="Arial", size=31, bold=True, color=NAVY)
    r2 = p.add_run("基于多智能体博弈推理的科研知识图谱发现系统")
    set_run_font(r2, cn=HEADING_FONT_CN, en="Arial", size=17, bold=True, color=TEAL_DARK)

    metadata = [
        ("推报学校名称", "中国矿业大学"),
        ("选题发榜单位", "上海云之脑智能科技有限公司"),
        ("竞榜选题名称", "领域知识个性化生成与多智能体协同决策系统研究比赛方案"),
        ("作品具体名称", "研海寻踪：基于多智能体博弈推理的科研知识图谱发现系统"),
        ("申报者姓名", "【待团队与导师审核后补充】"),
        ("指导教师", "【待补充】"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_geometry(table, [2300, 7060], indent_dxa=120)
    set_table_borders(table, color=WHITE, size=0)
    for label, value in metadata:
        cells = table.add_row().cells
        p1 = cells[0].paragraphs[0]
        style_paragraph(p1, alignment=WD_ALIGN_PARAGRAPH.RIGHT, after=3, line=1.2)
        r1 = p1.add_run(label + "：")
        set_run_font(r1, cn=HEADING_FONT_CN, en="Arial", size=11.5, bold=True, color=NAVY)
        p2 = cells[1].paragraphs[0]
        style_paragraph(p2, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=3, line=1.2)
        r2 = p2.add_run(value)
        set_run_font(r2, size=11.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("预审核初稿 · 仅供校内外专家内部把关\n2026 年 7 月")
    set_run_font(r, size=10.5, color=MUTED)


def add_color_overview(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("一页读懂“研海寻踪”")
    set_run_font(r, cn=HEADING_FONT_CN, en="Arial", size=25, bold=True, color=NAVY)
    p2 = doc.add_paragraph()
    style_paragraph(p2, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=16, line=1.2)
    r2 = p2.add_run("从论文证据到可信知识图谱，再到可验证的科研脉络与 Idea")
    set_run_font(r2, cn=HEADING_FONT_CN, en="Arial", size=13, bold=True, color=TEAL_DARK)

    cards = doc.add_table(rows=2, cols=2)
    cards.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(cards, [4680, 4680], indent_dxa=120)
    set_table_borders(cards, color=WHITE, size=4)
    card_data = [
        ("3 个垂直领域", "每个领域 30 篇可检索论文记录"),
        ("5 个协同角色", "2 个专职 Agent + 3 个核心裁决 Agent"),
        ("9 组完整样例", "领域 × 学习者画像，含完整输入—中间—输出"),
        ("84.6% 接收精确率", "24 条压力命题；UAR 由 63.6% 降至 18.2%"),
    ]
    for idx, (metric, detail) in enumerate(card_data):
        cell = cards.cell(idx // 2, idx % 2)
        set_cell_shading(cell, LIGHT_TEAL if idx != 3 else LIGHT_GOLD)
        p = cell.paragraphs[0]
        style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=3, line=1.15)
        r = p.add_run(metric + "\n")
        set_run_font(r, cn=HEADING_FONT_CN, en="Arial", size=18, bold=True, color=TEAL_DARK if idx != 3 else NAVY)
        r = p.add_run(detail)
        set_run_font(r, size=9.5, color=MUTED)

    add_note(
        doc,
        "核心闭环",
        "论文/证据卡 → 实体与关系候选 → 提出者—批判者—裁判 → accepted 图谱 → 意图驱动 GraphRAG → 技术脉络、待验证 Idea 与个性化学习资源。",
        fill=LIGHT_BLUE,
    )

    image_table = doc.add_table(rows=1, cols=2)
    image_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(image_table, [4680, 4680], indent_dxa=120)
    set_table_borders(image_table, color="D8E2DF", size=4)
    home = PROJECT_ROOT / "docs" / "assets" / "readme" / "demo-home.png"
    results = PROJECT_ROOT / "docs" / "assets" / "readme" / "demo-results.png"
    for cell, image_path, caption in (
        (image_table.cell(0, 0), home, "领域与学习者输入"),
        (image_table.cell(0, 1), results, "图谱、轨迹与消融输出"),
    ):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Cm(7.2))
        cp = cell.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = cp.add_run(caption)
        set_run_font(rr, size=9, color=MUTED)

    add_note(
        doc,
        "结果边界",
        "84.6% 来自开发可见的 24 条小样本压力集，不是公开科学信息抽取基准性能；“证据绑定 100%”只表示 accepted 关系具有有效 evidence_id，不代表引用语义完全正确。",
        fill="FCE9E5",
        color=RED,
    )


def add_abstract(doc: Document) -> None:
    add_heading(doc, "摘要", level=1, page_break=True)
    abstract = (
        "科研人员面对的困难已经从“找不到文献”转向“难以把跨论文、跨章节、跨学科的证据组织成可复核的知识”。"
        "传统关键词检索返回文献列表，通用大模型直接总结又容易丢失来源、混淆条件或生成未经证实的关联。"
        "围绕上海云之脑智能科技有限公司发布的“领域知识个性化生成与多智能体协同决策系统研究”榜题，"
        "本项目提出“研海寻踪”科研知识图谱发现系统：以论文结构化解析和证据跨度为底座，先由论文知识抽取 Agent "
        "形成实体、关系与主张候选，再通过“提出者—批判者—裁判”三个异质决策 Agent 分别完成候选组织、反证约束和独立裁决；"
        "用户意图感知 Agent 将检索型、分析型和 Idea 型查询分别路由至图广度、图深度和混合检索，"
        "最终输出可回指论文证据的技术脉络、争议提示、待验证研究假设及面向不同学习者的导读、实操和测评。"
        "当前 Demo 已构建科学信息抽取、材料发现与图神经网络、教育知识追踪三个垂直领域切片，每个领域收录 30 篇论文记录；"
        "其中 19 篇已形成证据卡，抽取 108 条实体证据跨度、66 个规范实体和 91 条候选关系。"
        "系统提供 3 组合成学习者画像及 9 组完整输入—协同中间数据—个性化输出样例。"
        "在固定 24 条候选命题的压力测试中，三智能体链路接收精确率为 84.6%（11/13），Gold 保留召回为 84.6%（11/13），"
        "不支持命题接收率为 18.2%（2/11）；相较最强单次判定基线，接收精确率提高 23.5 个百分点，不支持命题接收率下降 45.4 个百分点。"
        "系统同时实现本地离线运行、启动与请求超时、幂等、熔断、结构化日志和回归测试。"
        "上述结果证明了异质职责与证据约束在固定候选池上的工程价值，但尚不能替代大规模全文金标准、专家盲评和真实用户试点。"
        "项目下一阶段将接入 Docling/GROBID、GLiNER/GLiREL 或 OneKE，并在冻结的端到端测试集上完成实体、关系、证据跨度及下游发现效用评测。"
    )
    add_body(doc, abstract)
    p = doc.add_paragraph()
    style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=6, line=1.3)
    r1 = p.add_run("关键词：")
    set_run_font(r1, cn=HEADING_FONT_CN, en="Arial", size=11.5, bold=True, color=NAVY)
    r2 = p.add_run("科研知识图谱；科学信息抽取；证据溯源；多智能体协同；GraphRAG；个性化学习；研究发现")
    set_run_font(r2, size=11.5)


def add_toc(doc: Document) -> None:
    add_heading(doc, "目录", level=1, page_break=True)
    toc = [
        ("摘要", "3"),
        ("作品评估指标索引表", "5"),
        ("第一章 项目背景与榜单分析", "6"),
        ("第二章 国内外研究与竞品分析", "8"),
        ("第三章 研究问题与总体方案", "10"),
        ("第四章 核心方法：论文知识抽取与可信建图", "12"),
        ("第五章 意图驱动 GraphRAG 与个性化生成", "14"),
        ("第六章 系统设计、实现与使用", "16"),
        ("第七章 数据、实验与结果分析", "19"),
        ("第八章 创新点与技术指标实现", "21"),
        ("第九章 应用价值、竞品与实施计划", "23"),
        ("第十章 总结、局限与下一阶段", "25"),
        ("参考文献", "26"),
        ("附录", "27"),
    ]
    for title, page in toc:
        p = doc.add_paragraph()
        style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=4, line=1.2)
        p.paragraph_format.left_indent = Cm(0.2 if "第" in title else 0)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(14.3))
        r1 = p.add_run(title)
        set_run_font(r1, cn=HEADING_FONT_CN if title.startswith("第") else BODY_FONT_CN, en="Arial", size=11.5, bold=title.startswith("第"), color=NAVY if title.startswith("第") else INK)
        r2 = p.add_run("\t" + page)
        set_run_font(r2, size=11, color=MUTED)
    add_note(doc, "初稿说明", "目录页码按首版内容预算编排；提交定稿前需在 Word 中更新目录并逐页核对。")


def add_evaluation_index(doc: Document) -> None:
    add_heading(doc, "作品评估指标索引表", level=1, page_break=True)
    add_table(
        doc,
        ["评审维度", "本项目核心主张", "正文位置", "直接证据"],
        [
            ("作品完整性", "跑通论文/证据卡到图谱、检索、裁决和个性化资源的闭环", "第 3—6 章", "系统架构、接口输出、前端截图、运行说明"),
            ("技术性能", "固定候选池下，异质三智能体降低不支持命题入图", "第 7 章", "24 条压力集、四组对照、Wilson 区间、错误案例"),
            ("技术创新性", "证据优先建图、异质职责博弈、意图驱动图检索形成组合机制", "第 4、5、8 章", "数据契约、决策轨迹、消融矩阵"),
            ("赛题对齐", "3 个领域、3 组画像、9 组完整输入—中间—输出", "第 1、7、8 章", "领域注册表、完整样例 JSON、资源输出"),
            ("实用性", "支持论文推荐、机制分析、技术演化与待验证 Idea", "第 5、9 章", "GraphRAG 路由、图谱结果、场景流程"),
            ("工程可靠性", "离线可运行，具备超时、幂等、熔断、日志和测试", "第 6、7 章", "76 项自动化测试全部通过、六实验入口、启动脚本、健康检查"),
        ],
        [1500, 3360, 1500, 3000],
        font_size=9,
    )
    add_note(doc, "证据口径", "本表中的指标均对应当前提交版本。尚未形成的专利、软著、论文、应用证明和专家推荐不得作为已取得成果陈述。", fill="FCE9E5", color=RED)


def chapter_background(doc: Document) -> None:
    add_heading(doc, "第一章 项目背景与榜单分析", 1)
    add_heading(doc, "1.1 为什么“搜到论文”仍不足以支持科研创新", 2, page_break=False)
    add_body(
        doc,
        "科研文献规模持续增长，但研究者真正需要的通常不是另一份按相关度排序的列表，而是回答“某种方法为什么有效、在哪些数据与条件下有效、与哪些工作存在继承或冲突、下一步缺口在哪里”。这些答案往往分散在摘要、方法、实验表格、局限性和引文关系中。关键词检索以文档为基本单位，难以把跨论文证据组织成可计算关系；通用生成式问答虽然能够给出流畅总结，却可能省略限定条件、混淆论文贡献或生成不可追溯的关联。"
    )
    add_body(
        doc,
        "因此，本项目把问题重新定义为“可信的论文到知识图谱转换与图谱驱动发现”：任何进入正式图谱的关系都必须具有稳定实体端点、合法关系类型、可回指证据和明确裁决状态；任何由缺失边或低连接区域生成的 Idea 都必须标记为待验证假设，而非已证实结论。"
    )
    add_heading(doc, "1.2 榜单要求拆解", 2, page_break=False)
    add_table(
        doc,
        ["榜单要求", "本项目实现", "当前状态", "可复核位置"],
        [
            ("至少 1 个垂直领域知识库切片", "构建 3 个垂直领域，每个领域 30 篇检索记录", "Demo 级满足", "data/vertical_kb"),
            ("不少于 2 组差异化学习者数据", "本科科研入门、跨学科硕士、企业技术情报 3 组画像", "满足", "data/profiles"),
            ("多智能体协同中间数据", "专职 Agent 轨迹、三智能体轨迹、候选命题、批判意见、裁判分解", "满足", "agent_trace / claims"),
            ("最终个性化学习资源", "导读、复现实操、分阶测评、学习路径与回答骨架", "满足", "resources / report"),
            ("完整输入输出示例", "3 个领域 × 3 个画像 = 9 组固定样例", "满足", "complete_demo_cases.json"),
        ],
        [2100, 3300, 1450, 2510],
        font_size=8.8,
    )
    add_heading(doc, "1.3 研究目标与核心问题", 2, page_break=False)
    for item in (
        "RQ1：如何从论文结构和证据跨度中抽取可追溯的实体、关系与科研主张，而不是仅从题名或摘要生成裸三元组？",
        "RQ2：在候选池固定的条件下，提出者—批判者—裁判的异质职责是否能减少高置信错误和证据错配进入图谱？",
        "RQ3：如何根据检索、分析和 Idea 三类意图选择图广度、图深度或混合检索，从而兼顾覆盖、可解释路径和发现能力？",
        "RQ4：如何把同一知识底座转化为适配不同学情的导读、实操和测评，并保留完整决策过程？",
    ):
        add_bullet(doc, item)
    add_heading(doc, "1.4 项目定位", 2, page_break=False)
    add_note(
        doc,
        "定位",
        "“研海寻踪”不是通用聊天机器人，也不是单纯的图谱可视化页面；它是一套以证据可追溯为约束、以异质多智能体为决策机制、以科研知识图为推理底座的研究与工程原型。",
        fill=LIGHT_TEAL,
    )


def chapter_related(doc: Document) -> None:
    add_heading(doc, "第二章 国内外研究与竞品分析", 1)
    add_heading(doc, "2.1 科学信息抽取：从句内实体到全文关系", 2, page_break=False)
    add_body(
        doc,
        "DyGIE++ 以 span 表示联合建模实体、关系、事件和共指，SciREX 将任务扩展到全文显著实体与文档级 n 元关系，ReSel 进一步说明实验表格与正文必须联合建模。SciER 和 SciNLP 等近年工作则表明，科学实体、方法、数据集、任务与结果在跨句、跨章节和跨论文条件下仍具有显著抽取难度。由此，本项目把“章节—句子—字符跨度”写入证据对象，并把模型—数据集—指标—数值视为带条件的科研事件，而非拆散为无上下文边。"
    )
    add_heading(doc, "2.2 GraphRAG：从局部事实到全局主题", 2, page_break=False)
    add_body(
        doc,
        "Microsoft GraphRAG 通过实体、关系、社区和社区报告支持 Local、Global 与 DRIFT 等查询模式，为跨文档全局理解提供了成熟参照。本项目当前采用 GraphRAG-inspired 离线基线：用概念图、证据路径和连通社区实现广度、深度与混合检索，但不宣称已经运行微软官方 GraphRAG。下一阶段将通过 Bring Your Own Graph 接口导出 entities、relationships、text_units 与 communities，再比较官方查询模式的质量、成本和稳定性。"
    )
    add_heading(doc, "2.3 多智能体校验：角色数量不等于有效协作", 2, page_break=False)
    add_body(
        doc,
        "多智能体方法能够通过角色分工、辩论和审查提高复杂任务覆盖，但当角色共享相同模型、提示结构和证据时，也可能放大共同盲点。项目因此不以“Agent 数量”作为创新，而以职责异质性和证据依赖性作为可检验机制：提出者组织候选，批判者主动寻找类型、方向、强断言与证据问题，裁判根据独立规则作最终决策。当前同质三路投票与单次判定得到相同结果，正说明简单堆叠角色并不能自动带来收益。"
    )
    add_heading(doc, "2.4 产品与竞品能力矩阵", 2, page_break=False)
    add_table(
        doc,
        ["类型", "代表能力", "主要优势", "相对本项目的缺口"],
        [
            ("传统学术检索", "关键词、引用与排序", "覆盖大、使用成熟", "难以把跨论文证据转成可审计关系"),
            ("文献关系图工具", "引文/相似文献网络", "探索路径直观", "节点多为论文，缺少方法—任务—结果证据图"),
            ("通用 LLM/RAG", "自然语言问答与摘要", "交互门槛低", "主张、证据和裁决过程可能不透明"),
            ("企业技术情报平台", "专利、论文、产业数据聚合", "数据源丰富、流程成熟", "成本较高，个性化学习与开放实验有限"),
            ("研海寻踪", "证据跨度图 + 异质裁决 + 意图图检索", "每条关系可追溯，能展示失败与复核状态", "当前全文规模、盲评和真实用户试点仍不足"),
        ],
        [1550, 2300, 2350, 3160],
        font_size=8.8,
    )


def chapter_solution(doc: Document) -> None:
    add_heading(doc, "第三章 研究问题与总体方案", 1)
    add_heading(doc, "3.1 端到端技术路线", 2, page_break=False)
    add_body(
        doc,
        "系统输入包括领域论文或知识卡、学习者画像和研究查询。文献首先被转换为带章节和字符偏移的统一文档对象；论文知识抽取 Agent 输出实体、关系与证据候选；实体规范化模块合并别名并保留回滚记录；三智能体链路决定候选进入 accepted、needs_review 或 rejected 状态；用户意图感知 Agent 再选择图检索路线，最终生成论文推荐、机制分析、待验证 Idea 和个性化学习资源。"
    )
    arch = ASSET_DIR / "system-architecture.png"
    doc.add_picture(str(arch), width=Cm(15.3))
    add_caption(doc, "图 3-1 研海寻踪端到端系统架构")
    add_heading(doc, "3.2 统一口径：5 个协同角色，3 个核心决策 Agent", 2, page_break=False)
    add_table(
        doc,
        ["角色", "输入", "核心职责", "输出"],
        [
            ("论文知识抽取 Agent", "论文结构与证据片段", "抽取实体、关系、主张候选", "proposed 图谱与证据跨度"),
            ("用户意图感知 Agent", "研究查询", "识别检索/分析/Idea 意图并路由", "意图分数与图检索路线"),
            ("提出者 Agent", "抽取候选与 schema", "组织候选关系与依据", "候选命题、初始置信度"),
            ("批判者 Agent", "候选、端点与证据", "检查类型、方向、强断言和证据", "批判意见与风险标记"),
            ("裁判 Agent", "候选、批判、证据质量", "独立计算分数并做状态裁决", "accepted / needs_review / rejected"),
        ],
        [1800, 2200, 3100, 2260],
        font_size=8.8,
    )
    add_heading(doc, "3.3 决策目标与约束", 2, page_break=False)
    add_body(
        doc,
        "项目把最终入图视为选择性预测问题。裁判不仅追求接收命题数量，还需同时约束 accepted precision、Gold retention recall、Unsupported Acceptance Rate（UAR）和人工复核负担。没有证据、证据端点无效、关系越过 schema 或包含绝对化保证的候选应被拒绝或转入复核。系统允许拒答和保守，而不是以“回答所有问题”为目标。"
    )
    add_heading(doc, "3.4 最小可信知识单元", 2, page_break=False)
    add_note(
        doc,
        "数据契约",
        "最小可信单元 =（规范实体，关系，规范实体）+ paper_id + section/sentence/char span + extraction version + criticisms + judge score + status。裸三元组不能直接进入正式图谱。",
        fill=LIGHT_BLUE,
    )


def chapter_extraction(doc: Document) -> None:
    add_heading(doc, "第四章 核心方法：论文知识抽取与可信建图", 1)
    add_heading(doc, "4.1 文档结构解析", 2, page_break=False)
    add_body(
        doc,
        "当前 Demo 使用可复现的 PlainText/知识卡适配器，保留论文标识、标题、年份、章节、句子和字符偏移；对于正式全文路线，计划以 Docling 解析版面、阅读顺序、表格和公式，以 GROBID 的 TEI XML 强化作者、引文和参考文献链接。扫描件需启用 OCR 并记录 OCR 置信度。解析器版本、schema 版本和原文件哈希被视为图谱来源的一部分。"
    )
    add_heading(doc, "4.2 领域本体与 schema 约束", 2, page_break=False)
    add_table(
        doc,
        ["实体类型", "示例", "关系类型", "约束示例"],
        [
            ("METHOD", "GraphRAG、GLiNER、CGCNN", "USES / EXTENDS", "METHOD → METHOD"),
            ("TASK", "文档级关系抽取、知识追踪", "ADDRESSES", "METHOD → TASK"),
            ("DATASET", "SciREX、EdNet", "EVALUATES_ON", "METHOD → DATASET"),
            ("METRIC", "F1、证据命中率", "REPORTS", "实验/数据集 → METRIC"),
            ("FINDING / LIMITATION", "OOD 下降、标注噪声", "SUPPORTS / CONTRADICTS", "主张 ↔ 发现"),
            ("DOMAIN", "NLP、材料、教育", "RELATED_TO", "默认 needs_review"),
        ],
        [1600, 2700, 1900, 3160],
        font_size=8.8,
    )
    add_heading(doc, "4.3 候选实体、关系与证据跨度", 2, page_break=False)
    add_body(
        doc,
        "当前候选器由版本化词典、别名规则和触发模式构成，优势是可解释、无外部模型依赖和现场稳定；它是工程下限，不代表最终模型。下一阶段采用三路候选并集：高精度规则路、SciBERT/DeepKE/OneKE 监督路、GLiNER/GLiREL 或结构化 LLM 低资源路。所有候选必须绑定章节、句子、字符起止位置和 paper_id。未取得全文的元数据记录只能参与书目检索，代码会将其排除在证据图和关系抽取之外。"
    )
    add_heading(doc, "4.4 实体规范化与跨论文融合", 2, page_break=False)
    add_body(
        doc,
        "实体融合依次执行 Unicode 与大小写规范化、连字符处理、缩写展开、别名精确匹配和类型兼容检查。正式版本将增加向量召回与交叉编码器重排，并引入邻接结构、作者、时间和版本特征。低于阈值的候选不会被强制合并，而是保留为独立实体进入人工复核；合并操作必须可撤销并记录前后值。"
    )
    add_heading(doc, "4.5 提出者—批判者—裁判的博弈式校验", 2, page_break=False)
    add_numbered_list(
        doc,
        (
            "提出者读取抽取图谱和 schema，形成带来源、类型、候选置信度和 proposal_reason 的命题。",
            "批判者验证 evidence_id、证据语义、实体端点、类型方向、绝对化措辞和单来源风险，并写入 criticisms。",
            "裁判结合基础分、证据奖励、交叉来源奖励与风险惩罚生成 judge_score，输出 accepted、needs_review 或 rejected。",
            "只有 accepted 关系进入下游资源；needs_review 保留给人工审查，rejected 不进入正式概念图。",
        ),
    )
    add_heading(doc, "4.6 图融合、社区与增量状态", 2, page_break=False)
    add_body(
        doc,
        "关系状态按 raw → proposed → reviewed → accepted 迁移。重复关系聚合证据但不覆盖来源，互相冲突的发现可通过 CONTRADICTS 并存。图谱保存首次出现、最近更新、论文发表时间、数据/解析/schema 版本和自动/人工决策。当前以连通分量形成可复现社区，正式版将比较 Leiden 分层社区和社区摘要。"
    )
    add_heading(doc, "4.7 当前证据图规模", 2, page_break=False)
    add_table(
        doc,
        ["领域", "收录论文", "证据卡", "证据跨度", "规范实体", "候选关系"],
        [
            ("科学信息抽取与知识图谱", "30", "8", "47", "31", "40"),
            ("材料发现与图神经网络", "30", "5", "28", "17", "24"),
            ("教育知识追踪与个性化学习", "30", "6", "33", "18", "27"),
            ("合计", "90", "19", "108", "66", "91"),
        ],
        [2800, 1200, 1200, 1400, 1300, 1460],
        font_size=9,
    )
    add_note(
        doc,
        "重要边界",
        "91 条是进入裁决流程的候选关系，不等于 91 条均被领域专家判真；100% 证据绑定只表示候选存在有效本地证据跨度，不等于关系正确率为 100%。",
        fill="FCE9E5",
        color=RED,
    )


def chapter_graphrag(doc: Document) -> None:
    add_heading(doc, "第五章 意图驱动 GraphRAG 与个性化生成", 1, page_break=False)
    add_heading(doc, "5.1 为什么检索、分析和 Idea 不能共用同一路线", 2, page_break=False)
    add_body(
        doc,
        "信息检索强调覆盖与多样性，机制分析强调少量可信多跳路径，Idea 发现则需要从全局社区进入局部关系并识别缺失边。统一检索器容易在三类目标之间折中失衡，因此系统先识别用户意图，再选择图算法和上下文预算。"
    )
    add_table(
        doc,
        ["意图", "路由", "核心策略", "主要输出"],
        [
            ("文献检索", "graph_breadth", "多种子 BFS 展开两层，按证据聚合论文", "推荐论文、命中概念、覆盖范围"),
            ("分析推理", "graph_depth", "锚定实体，搜索最多三跳可信路径", "可审计 triples、relation IDs、evidence IDs"),
            ("Idea 发现", "hybrid_drift", "社区 primer + 局部深挖 + 缺失边检测", "待验证假设、证据来源、后续问题"),
        ],
        [1500, 1700, 3600, 2560],
        font_size=9,
    )
    add_heading(doc, "5.2 用户意图感知 Agent", 2, page_break=False)
    add_body(
        doc,
        "当前版本采用可解释的多标签词法路由，输出 primary_intent、secondary_intents、confidence、matched_signals 和 score_breakdown。其目的不是假装使用复杂模型，而是建立可测契约。下一阶段可使用 Qwen2.5-3B-Instruct 结构化分类或 bge-m3 向量路由，并在独立标注意图集上报告 macro-F1、route accuracy、ECE 和 fallback rate。"
    )
    add_heading(doc, "5.3 图谱辅助理解技术脉络", 2, page_break=False)
    add_body(
        doc,
        "系统按照论文年份、实体关系和来源证据生成技术演化时间线。例如，科学信息抽取领域可把 span 联合抽取、全文文档级关系、文本—表格联合抽取、开放类型实体/关系抽取组织为带论文依据的路径。图谱不只显示“论文引用论文”，而是显示“方法解决任务、方法使用机制、方法在哪个数据集上评价、论文报告什么结果”。"
    )
    add_heading(doc, "5.4 图谱驱动 Idea", 2, page_break=False)
    add_body(
        doc,
        "Idea 模块从低交叉验证关系、缺失 EVALUATES_ON 边、跨领域弱连接和争议关系形成候选假设。每条 Idea 保存生成路径与 evidence_ids，并强制标注 novelty_status=unverified。系统只能说明“现有切片中缺少证据”或“值得进一步查证”，不能据此证明全球范围内的新颖性。正式版必须联网检索、查新并由专家评价可行性与新颖性。"
    )
    add_heading(doc, "5.5 个性化资源生成", 2, page_break=False)
    add_body(
        doc,
        "学习者画像包含目标、兴趣、知识分数、偏好、所需概念和预期难度。诊断服务计算准备度、盲点和目标难度，资源 Agent 只读取 accepted 关系和画像命中的证据，生成定制导读、复现实操与分阶测评。反馈“太难/合适/太简单”会调整下一轮目标难度和解释策略。当前画像为合成数据，避免处理真实学生隐私。"
    )


def chapter_system(doc: Document) -> None:
    add_heading(doc, "第六章 系统设计、实现与使用", 1)
    add_heading(doc, "6.1 技术栈", 2, page_break=False)
    add_table(
        doc,
        ["层次", "当前实现", "选择理由"],
        [
            ("前端", "HTML5 + CSS3 + 原生 JavaScript + SVG", "零构建依赖，适合离线比赛演示"),
            ("后端", "Python 3.11+ 标准库 HTTP 服务", "轻量、可打包、单进程稳定"),
            ("知识加工", "版本化 schema、规则抽取、实体融合、证据审计", "形成可运行与可解释下限"),
            ("图检索", "GraphRAG-inspired BFS/多跳/混合路由", "展示不同意图的图检索差异"),
            ("存储", "JSON + SQLite 导出", "便于审计、复现和演示"),
            ("工程 Harness", "超时、幂等、熔断、有界并发、日志、健康探针", "避免网络/并发导致 Demo 卡死"),
            ("交付", "Windows 双击脚本 + Docker Compose", "接收者可本地复现"),
        ],
        [1600, 3600, 4160],
        font_size=9,
    )
    add_heading(doc, "6.2 后端可靠性设计", 2, page_break=False)
    for item in (
        "启动层：后台进程不长期占用调用端句柄，总启动截止默认 10 秒；健康请求单次最多等待 1 秒。",
        "HTTP 层：限制请求体和字符编码，输出稳定错误结构；写请求使用 Idempotency-Key 防止重复运行。",
        "任务层：有界 worker 与队列，过载返回 429，任务超时返回 504。",
        "联网 RAG：有限重试、指数退避和 closed/open/half-open 熔断；失败时回退本地缓存。",
        "可观测性：请求 ID、运行 ID、结构化日志、聚合指标和 append-only 运行日志。",
        "安全边界：默认只绑定 127.0.0.1；非回环部署无 Token 时拒绝启动。",
    ):
        add_bullet(doc, item)
    add_heading(doc, "6.3 前端界面与演示路径", 2, page_break=False)
    doc.add_picture(str(PROJECT_ROOT / "docs" / "assets" / "readme" / "demo-home.png"), width=Cm(15.2))
    add_caption(doc, "图 6-1 领域知识库与学习者画像选择界面")
    doc.add_picture(str(PROJECT_ROOT / "docs" / "assets" / "readme" / "demo-results.png"), width=Cm(15.2))
    add_caption(doc, "图 6-2 协同轨迹、图谱检索与消融结果界面")
    add_heading(doc, "6.4 使用方式", 2, page_break=False)
    add_numbered_list(
        doc,
        (
            "完整解压提交 ZIP，安装 Python 3.11 或更高版本并勾选 Add Python to PATH。",
            "双击 RUN_DEMO.bat；健康检查通过后访问 http://127.0.0.1:8765/。",
            "选择领域、学习者画像和研究任务，点击“启动协同推理”。",
            "依次查看知识抽取统计、意图路由、三智能体轨迹、证据图谱、消融与个性化资源。",
            "演示结束双击 STOP_DEMO.bat；若浏览器未自动打开，双击 OPEN_DEMO.url。",
        ),
    )
    add_note(doc, "离线能力", "除 OpenAlex 联网扩展外，领域切片、建图、三智能体裁决、GraphRAG 路由、消融和个性化资源均可离线运行。")


def chapter_experiments(doc: Document) -> None:
    add_heading(doc, "第七章 数据、实验与结果分析", 1)
    add_heading(doc, "7.1 数据构成与来源边界", 2, page_break=False)
    add_body(
        doc,
        "三个领域采用目的性分层选择：科学信息抽取验证项目核心方法，材料发现验证跨学科图结构，教育知识追踪直接对齐榜题中的学习者画像与资源生成。每个领域 30 篇记录由 DOI、官方来源和 Crossref 候选快照支持；其中 19 篇证据卡用于建图，其余 71 篇 metadata_only 记录只参与书目检索。该三领域集合不是统计意义上的随机样本，不能推断系统在全部学科上的平均性能。"
    )
    add_heading(doc, "7.2 两条评测轨道", 2, page_break=False)
    add_body(
        doc,
        "Track A 固定 24 条候选命题，只改变决策机制，回答异质批判与裁决是否能减少错误入图。压力集包含 13 条支持命题和 11 条不支持命题，以及低置信真阳性、有效证据 ID 语义错配、缺失/错误证据、绝对化和类型方向错误。Track B 是后续必须完成的端到端全文抽取盲测，将比较规则、GLiNER/GLiREL、SciBERT/DyGIE++、DeepKE/OneKE、单次 LLM 和最佳候选器 + 三智能体。"
    )
    add_heading(doc, "7.3 对照结果", 2, page_break=False)
    ablation = ASSET_DIR / "ablation.png"
    doc.add_picture(str(ablation), width=Cm(15.2))
    add_caption(doc, "图 7-1 固定候选池下四种决策机制的接收精确率与 UAR")
    add_table(
        doc,
        ["变体", "接收精确率", "Gold 召回", "F1", "UAR", "接收/复核/拒绝"],
        [
            ("E0 普通规则程序", "50.0%（11/22）", "84.6%（11/13）", "62.9%", "100.0%（11/11）", "22 / 0 / 2"),
            ("E1 单次判定", "61.1%（11/18）", "84.6%（11/13）", "71.0%", "63.6%（7/11）", "18 / 0 / 6"),
            ("E2 同质三路投票", "61.1%（11/18）", "84.6%（11/13）", "71.0%", "63.6%（7/11）", "18 / 0 / 6"),
            ("E3 提出—批判—裁判", "84.6%（11/13）", "84.6%（11/13）", "84.6%", "18.2%（2/11）", "13 / 2 / 9"),
        ],
        [2200, 1650, 1550, 1000, 1600, 1360],
        font_size=8.5,
    )
    add_body(
        doc,
        "相较最强基线 E1/E2，E3 接收精确率提高 23.5 个百分点，UAR 下降 45.4 个百分点，Gold 召回保持不变。E3 接收精确率和 Gold 召回的 Wilson 95% 区间均为 [0.578, 0.957]，UAR 区间为 [0.051, 0.477]。由于样本规模较小且开发者知道错误类型，结果应解释为“机制可行性与回归证据”，不能表述为统计上显著的公开基准领先。"
    )
    add_heading(doc, "7.4 失败案例与误差分析", 2, page_break=False)
    add_table(
        doc,
        ["案例", "现象", "原因", "后续修复方向"],
        [
            ("B019/B020", "E3 仍接收两条不支持命题", "evidence_id 有效但语义不支持，结构护栏无法识别", "引入 NLI/claim verification 与句级证据充分性模型"),
            ("B013/B014", "两条支持命题进入复核而非接收", "基础置信度较低，裁判偏保守", "验证集校准、选择性预测与人工复核成本优化"),
            ("元数据论文", "不能生成关系", "未获得并解析全文", "合法取得全文后重新解析与双人标注"),
            ("Idea 新颖性", "只能提示切片中的缺失边", "图谱覆盖有限", "联网查新、专利检索和专家盲评"),
        ],
        [1400, 2350, 2700, 2910],
        font_size=8.8,
    )
    add_heading(doc, "7.5 自动化与工程验收", 2, page_break=False)
    add_body(
        doc,
        "截至本初稿生成前，Python 源码编译、Node 前端语法检查和完整 unittest 均通过。测试发现 76 项，全部属于当前验收或版本边界检查并通过，跳过数与失败数均为 0。测试覆盖领域注册、每领域 30 篇记录、证据端点、关系绑定、意图路由、多跳路径、三智能体版本边界、六组实验入口、幂等、熔断、超时、部署路径和前端 API 契约。"
    )
    add_note(
        doc,
        "指标解释",
        "页面中的单轮“幻觉代理率 0%、适配准确度 100%、知识覆盖率 100%”是结构性健康检查或画像覆盖代理，不作为作品效果主指标。作品书统一使用带分子/分母和失败案例的 24 条压力集结果。",
        fill="FCE9E5",
        color=RED,
    )


def chapter_innovations(doc: Document) -> None:
    add_heading(doc, "第八章 创新点与技术指标实现", 1)
    add_heading(doc, "8.1 创新点一：证据跨度驱动的论文知识图谱", 2, page_break=False)
    add_body(
        doc,
        "相较只对标题摘要生成三元组的常见 Demo，本项目把章节、句子、字符跨度、schema、实体规范化和关系状态纳入统一数据契约。每条关系能够沿着“关系—实体—证据—论文”链路回查，元数据论文被强制排除在关系抽取之外。该创新的实质是把可追溯性从展示功能提升为入图前置约束。"
    )
    add_heading(doc, "8.2 创新点二：异质三智能体的选择性裁决", 2, page_break=False)
    add_body(
        doc,
        "系统将候选生成、反证审查和最终决策分离，允许复核与拒答；同质三路投票作为反例被纳入对照。24 条压力集显示，收益来自职责和检查信号的差异，而不是 Agent 数量。下一阶段将进一步比较共享模型与异构模型、全量辩论与选择性辩论的质量、时延和 token 成本。"
    )
    add_heading(doc, "8.3 创新点三：意图驱动的图检索与科研发现", 2, page_break=False)
    add_body(
        doc,
        "系统把“找论文、做分析、想 Idea”映射为广度、深度和混合图检索，让不同查询使用不同路径预算和输出契约。Idea 由图谱缺失边和弱连接产生，并带来源与未验证标记，避免把生成式灵感包装为事实。"
    )
    add_heading(doc, "8.4 创新点四：知识底座与学情画像联合决策", 2, page_break=False)
    add_body(
        doc,
        "同一 accepted 图谱针对本科科研入门、跨学科硕士和企业技术情报画像生成不同的解释深度、学习路径、实操任务和测评。画像不改变事实层，只改变检索重点和资源编排，从架构上避免“为了个性化而改写事实”。"
    )
    add_heading(doc, "8.5 技术指标完成表", 2, page_break=False)
    add_table(
        doc,
        ["指标", "目标", "当前结果", "结论"],
        [
            ("垂直知识库", "≥1 个领域", "3 个领域，每领域 30 篇", "Demo 级完成"),
            ("学习者画像", "≥2 组", "3 组合成画像", "完成"),
            ("完整案例", "含输入/中间/输出", "9 组机器可读案例", "完成"),
            ("可信裁决", "降低错误入图", "precision 84.6%，UAR 18.2%", "Track A 初步支持"),
            ("证据可追溯", "关系绑定来源", "91 条候选均有结构证据跨度", "结构护栏完成"),
            ("工程可复现", "可离线运行并自检", "76 项测试通过，0 跳过、0 失败；六实验入口可执行", "完成"),
            ("端到端抽取主实验", "全文冻结测试集", "尚未完成", "下一阶段核心"),
            ("知识产权/应用证明", "按校赛要求补强", "待学校流程与导师审核", "不得提前宣称"),
        ],
        [1900, 2300, 3000, 2160],
        font_size=8.7,
    )


def chapter_application(doc: Document) -> None:
    add_heading(doc, "第九章 应用价值、竞品与实施计划", 1)
    add_heading(doc, "9.1 高校科研场景", 2, page_break=False)
    add_body(
        doc,
        "面向本科生和研究生，系统可将一个陌生领域拆解为核心方法、任务、数据集、指标和证据路径，降低阅读门槛；对导师和课题组，可用于构建实验室领域图谱、追踪技术演化和检查学生综述中的证据来源。系统生成的 Idea 只作为讨论起点，不能替代查新、实验与导师判断。"
    )
    add_heading(doc, "9.2 企业技术情报场景", 2, page_break=False)
    add_body(
        doc,
        "面向研发与战略部门，系统可围绕某项技术形成证据图、关键论文、能力边界和待验证路线，帮助团队从“人工阅读大量文献”转向“先看关系与证据，再回读重点论文”。企业版本需要增加账号权限、私有语料隔离、任务队列、审计和私有化部署。"
    )
    add_heading(doc, "9.3 个性化科研训练场景", 2, page_break=False)
    add_body(
        doc,
        "榜题强调学习者画像与多智能体决策。本项目把图谱证据与画像结合，生成不同难度的导读、实操和测评，可用于科研方法课、文献研读训练和企业内部技术培训。未来需要在真实学习者上开展前后测、任务完成时间和主观有用性评估。"
    )
    add_heading(doc, "9.4 实施路线", 2, page_break=False)
    add_table(
        doc,
        ["阶段", "核心任务", "验收证据"],
        [
            ("M1：全文 Pilot", "30 篇合法全文；10 篇双人标注；固化 schema", "解析质量、标注一致性、首版金标准"),
            ("M2：候选模型", "GLiNER/GLiREL、OneKE、SciBERT 基线", "实体/关系/证据跨度 F1"),
            ("M3：端到端盲测", "冻结 60+ 关系、跨领域/OOD 与表格子集", "主结果、消融、置信区间、错误分析"),
            ("M4：GraphRAG 接入", "BYOG、Leiden、社区报告与 embeddings", "质量/时延/token 成本对比"),
            ("M5：真实试点", "高校课题组与企业情报用户任务", "完成时间、有用性、复核负担与应用证明"),
        ],
        [1700, 4300, 3360],
        font_size=8.8,
    )
    add_heading(doc, "9.5 风险与合规", 2, page_break=False)
    for item in (
        "论文版权：无法再分发的全文只保存来源链接、哈希和合法标注偏移，不将论文 PDF 打入公开仓库。",
        "事实风险：所有生成回答必须引用 accepted 图谱；Idea 明确标注待验证，保留人工复核入口。",
        "隐私风险：当前使用合成画像；真实学习数据接入前需最小化采集、授权、脱敏和访问控制。",
        "模型与开源许可：第三方组件须记录版本、许可证和修改范围；自研代码与第三方代码边界清晰。",
        "生产风险：当前是单机竞赛 Demo，不等同于多租户生产系统；公网长期部署前需 TLS、认证、共享状态、监控和备份。",
    ):
        add_bullet(doc, item)


def chapter_conclusion(doc: Document) -> None:
    add_heading(doc, "第十章 总结、局限与下一阶段", 1)
    add_heading(doc, "10.1 项目总结", 2, page_break=False)
    add_body(
        doc,
        "“研海寻踪”围绕科研文献碎片化与榜题中的个性化、多智能体协同要求，完成了从领域文献切片、证据跨度抽取、可信建图、异质三智能体裁决到意图驱动图检索和个性化资源生成的可运行闭环。项目的核心价值不在于页面展示，而在于把“每条关系能否回到证据、错误能否被批判、裁判为何接收或拒绝”变成显式数据与可复现实验。"
    )
    add_heading(doc, "10.2 当前局限", 2, page_break=False)
    for item in (
        "19 篇证据卡仍以项目组释义和公开摘要为主，尚未形成大规模合法全文解析集。",
        "24 条压力命题样本较小且开发可见，存在机制过拟合风险。",
        "候选器以规则为主，尚未完成 GLiNER/GLiREL、OneKE 等模型主实验。",
        "未完成双人独立标注、第三人仲裁、专家盲评和真实学习者试点。",
        "未取得可在本初稿中核验的学生近两年论文、专利、软著或应用证明。",
    ):
        add_bullet(doc, item)
    add_heading(doc, "10.3 下一阶段优先级", 2, page_break=False)
    add_numbered_list(
        doc,
        (
            "优先完成全文 Pilot、标注手册和金标准，先把论文信息抽取做实。",
            "接入至少两类候选模型，冻结 Track B 测试集并完成端到端主结果与消融。",
            "将当前图导出为 GraphRAG BYOG 数据，比较 Local/Global/DRIFT 与离线基线。",
            "开展 2—3 组真实用户任务，获取导师、科研人员或企业分析师的可核验反馈。",
            "在学校指导下推进软著/专利/论文，所有权属成果在获得受理或证书后再写入定稿。",
        ),
    )


def add_references(doc: Document) -> None:
    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Edge D, Trinh H, Cheng N, et al. From Local to Global: A Graph RAG Approach to Query-Focused Summarization. arXiv:2404.16130, 2024.",
        "[2] Luan Y, Wadden D, He L, et al. A General Framework for Information Extraction using Dynamic Span Graphs. NAACL, 2019.",
        "[3] Wadden D, Wennberg U, Luan Y, et al. Entity, Relation, and Event Extraction with Contextualized Span Representations. EMNLP-IJCNLP, 2019.",
        "[4] Jain S, van Zuylen M, Hajishirzi H, et al. SciREX: A Challenge Dataset for Document-Level Information Extraction. ACL, 2020.",
        "[5] Viswanathan V, Neubig G, et al. ReSel: N-ary Relation Extraction from Scientific Text and Tables. EMNLP, 2022.",
        "[6] Veyseh A P B, Dernoncourt F, Nguyen T H. GLiNER: Generalist Model for Named Entity Recognition using Bidirectional Transformer. NAACL, 2024.",
        "[7] Zaratiana U, Tomeh N, Holat P, et al. GLiREL: Generalist Model for Zero-Shot Relation Extraction. NAACL, 2025.",
        "[8] Lo K, Wang L L, Neumann M, et al. S2ORC: The Semantic Scholar Open Research Corpus. ACL, 2020.",
        "[9] Auer C, Dolfi M, Carvalho A, et al. Docling Technical Report. arXiv:2408.09869, 2024.",
        "[10] Lairgi Y, Moncla L, Cazabet R, et al. iText2KG: Incremental Knowledge Graphs Construction Using Large Language Models. arXiv:2409.03284, 2024.",
        "[11] Li G, Hammoud H A A K, Itani H, et al. CAMEL: Communicative Agents for Mind Exploration of Large Scale Language Model Society. NeurIPS, 2023.",
        "[12] Wu Q, Bansal G, Zhang J, et al. AutoGen: Enabling Next-Gen LLM Applications via Multi-Agent Conversation. arXiv:2308.08155, 2023.",
        "[13] Madaan A, Tandon N, Gupta P, et al. Self-Refine: Iterative Refinement with Self-Feedback. NeurIPS, 2023.",
        "[14] Piech C, Bassen J, Huang J, et al. Deep Knowledge Tracing. NeurIPS, 2015.",
        "[15] Merchant A, Batzner S, Schoenholz S S, et al. Scaling Deep Learning for Materials Discovery. Nature, 2023.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=4, line=1.2)
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)
        run = p.add_run(ref)
        set_run_font(run, size=9.5, color=INK)


def add_appendix(doc: Document) -> None:
    add_heading(doc, "附录 A：Demo 复核清单", 1)
    for item in (
        "解压 ZIP 后确认根目录存在 RUN_DEMO.bat、STOP_DEMO.bat、OPEN_DEMO.url、GITHUB_REPOSITORY.url。",
        "双击运行，检查 /api/health 和 /api/ready 正常，页面显示 3 个领域和 90 篇论文记录。",
        "运行任意领域与画像，检查 2 个专职 Agent、3 个核心决策 Agent 及 3 项服务轨迹。",
        "核对知识图谱中的论文—证据—实体链、裁判分解和 24 条压力集消融。",
        "点击 GitHub 快捷方式，核对默认 main 分支与提交历史。",
    ):
        add_bullet(doc, item)
    add_heading(doc, "附录 B：当前支撑材料", 1)
    add_table(
        doc,
        ["材料", "位置/状态", "用途"],
        [
            ("源代码仓库", "GitHub main，提交 63f6078", "复核代码与历史"),
            ("Windows Demo ZIP", "dist/yanhai-demo-windows.zip", "离线运行"),
            ("领域知识库", "data/vertical_kb", "3 领域 × 30 篇"),
            ("完整样例", "data/examples/complete_demo_cases.json", "输入—中间—输出"),
            ("压力测试集", "data/evaluation/decision_benchmark.json", "四组对照与错误案例"),
            ("架构、路线与实验文档", "docs/08、11、14、15、16", "技术与复现说明"),
            ("参考作品书", "references/competition_examples", "格式与结构参考"),
        ],
        [2100, 3600, 3660],
        font_size=8.8,
    )
    add_heading(doc, "附录 C：定稿前必须补充", 1)
    add_note(
        doc,
        "导师审核项",
        "申报者与指导教师姓名、正式封面届次、联系方式、学校/学院标识、团队分工与贡献说明、经导师确认的创新表述。",
        fill=LIGHT_GOLD,
    )
    add_note(
        doc,
        "支撑材料缺口",
        "学生近两年论文、专利、软著、应用证明、专家推荐等目前无可核验材料。若团队已有，请只在取得证书、受理通知、正式论文或盖章证明后加入，并建立与创新点的对应关系。",
        fill="FCE9E5",
        color=RED,
    )
    add_note(
        doc,
        "实验缺口",
        "合法全文、双人标注与仲裁、端到端 Track B、真实用户试点和专家盲评是作品从“完整 Demo”提升为“优质科研作品”的决定性工作。",
        fill=LIGHT_BLUE,
    )


def build_document(output_path: Path) -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    build_architecture_figure(ASSET_DIR / "system-architecture.png")
    build_ablation_figure(ASSET_DIR / "ablation.png")

    doc = Document()
    configure_styles(doc)
    add_cover(doc)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_furniture(body_section, first_page=False)
    add_color_overview(doc)
    add_abstract(doc)
    add_toc(doc)
    add_evaluation_index(doc)
    chapter_background(doc)
    chapter_related(doc)
    chapter_solution(doc)
    chapter_extraction(doc)
    chapter_graphrag(doc)
    chapter_system(doc)
    chapter_experiments(doc)
    chapter_innovations(doc)
    chapter_application(doc)
    chapter_conclusion(doc)
    add_references(doc)
    add_appendix(doc)

    core = doc.core_properties
    core.title = "研海寻踪：基于多智能体博弈推理的科研知识图谱发现系统"
    core.subject = "挑战杯揭榜挂帅专项赛作品书预审核初稿"
    core.author = "研海寻踪项目组"
    core.keywords = "科研知识图谱, 多智能体, GraphRAG, 科学信息抽取, 证据溯源"
    core.comments = "预审核初稿；提交前须经指导老师确认。"
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, default=OUTPUT_DIR / DOCX_NAME)
    args = parser.parse_args()
    build_document(args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
